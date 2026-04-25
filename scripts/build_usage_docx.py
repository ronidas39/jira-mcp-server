"""Generate `docs/USAGE.docx`: the end-to-end user guide for this repo.

The text lives in this script rather than in a markdown source plus a
pandoc step so the file is reproducible with nothing on disk besides
python-docx (already a dev dep transitive). Re-run after substantive
changes to keep the guide and the code in sync:

    python scripts/build_usage_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "docs" / "USAGE.docx"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _code_block(doc: Document, text: str) -> None:
    """Render a block of code in a single shaded, monospace paragraph."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.width = Cm(16)
    _shade(cell, "F4F4F4")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def _inline_code(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _kv_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for col, label in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            table.rows[r].cells[col].text = value


def _bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)


def _subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _section_break(doc: Document) -> None:
    doc.add_paragraph().add_run("").add_break()


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def _para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    _title(doc, "Jira MCP Server")
    _subtitle(doc, "End-to-End Usage Guide")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Backend, frontend, and operator workflows for the Model Context "
        "Protocol server that exposes Atlassian Jira Cloud as tools, "
        "resources, and prompts."
    )
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Repository: https://github.com/ronidas39/jira-mcp-server")
    run.font.size = Pt(10)
    run.italic = True
    doc.add_page_break()

    # 1. What this is
    _h1(doc, "1. What this is")
    _para(
        doc,
        "This project ships two halves of one product. The first is a "
        "Python MCP (Model Context Protocol) server that talks to Atlassian "
        "Jira Cloud and exposes the day-to-day operations a team runs (issue "
        "CRUD, JQL search, sprint reporting, workflow transitions, comments, "
        "and a small analytics surface) as MCP tools, resources, and prompt "
        "templates. The second is a custom Next.js dashboard that consumes "
        "those tools so a human operator can drive Jira from a browser "
        "without ever opening Atlassian's own UI.",
    )
    _para(
        doc,
        "Audit logging (every write produces an immutable Mongo row), "
        "structured JSON logs, retry on 429 and 5xx with exponential "
        "backoff, and HTTPS-only validation are baked in. The same server "
        "runs locally over stdio for IDE assistants and over streamable "
        "HTTP for the bundled web UI or any remote MCP client.",
    )

    # 2. Architecture
    _h1(doc, "2. System architecture")
    _para(
        doc,
        "Three processes in a typical deployment: the Python MCP server, "
        "MongoDB (audit log plus TTL cache), and the Next.js web UI. The "
        "web UI talks to the MCP server over streamable HTTP through a "
        "thin set of API routes; the MCP server talks to Jira Cloud over "
        "HTTPS and to MongoDB over the wire protocol.",
    )
    _code_block(
        doc,
        "  Browser  --HTTP-->  Next.js (web/)  --HTTP-->  MCP server  --HTTPS-->  Jira Cloud\n"
        "                                                       |\n"
        "                                                       +--TCP--> MongoDB\n",
    )
    _para(
        doc,
        "The web UI also has an optional /chat page that uses the Anthropic "
        "API; it does not bypass the MCP server, it just lets the model "
        "drive the same MCP tools the rest of the UI exposes through forms.",
    )

    # 3. Prerequisites
    _h1(doc, "3. Prerequisites")
    _kv_table(
        doc,
        ["Component", "Minimum version", "Notes"],
        [
            ["Python", "3.11", "3.12 is what CI runs against."],
            ["Node.js", "20.x", "Tailwind v4 and Next.js 15 require recent Node."],
            ["MongoDB", "6.0", "Local Docker container or MongoDB Atlas."],
            ["Atlassian account", "Cloud only", "API token or OAuth 3LO app."],
            ["Docker (optional)", "24.x", "Only needed for the bundled compose stack."],
        ],
    )
    _para(
        doc,
        "Operating system: developed and tested on macOS and Linux. The "
        "Python.org build of Python 3.13 on macOS needs the certifi CA "
        "bundle for Atlas TLS; the Mongo connection helper handles that "
        "automatically.",
    )

    # 4. Quick start
    _h1(doc, "4. Quick start")
    _para(doc, "The shortest path from a fresh clone to a running stack:")
    _code_block(
        doc,
        "git clone https://github.com/ronidas39/jira-mcp-server.git\n"
        "cd jira-mcp-server\n"
        "\n"
        "# 1. Python deps and config\n"
        "pip install -e \".[dev]\"\n"
        "cp .env.example .env\n"
        "# edit .env: JIRA_BASE_URL, JIRA_EMAIL, JIRA_API_TOKEN, MONGO_URI\n"
        "\n"
        "# 2. Local MongoDB plus indexes\n"
        "docker compose up -d mongodb\n"
        "python scripts/init_db.py\n"
        "\n"
        "# 3. Smoke against your real Jira\n"
        "PYTHONPATH=src python scripts/smoke_jira.py\n"
        "\n"
        "# 4. Start the MCP server over HTTP (terminal 1)\n"
        "MCP_TRANSPORT=http MCP_HTTP_PORT=8765 PYTHONPATH=src python -m jira_mcp\n"
        "\n"
        "# 5. Start the UI (terminal 2)\n"
        "cd web\n"
        "npm install\n"
        "cp .env.example .env.local   # set MCP_SERVER_URL=http://localhost:8765/mcp\n"
        "npm run dev                  # http://localhost:3000\n",
    )
    _para(
        doc,
        "If Docker is your preferred path, jump to section 9; the compose "
        "stack brings Mongo, the server, and the UI online together.",
    )

    # 5. Backend setup
    _h1(doc, "5. Backend setup (Python MCP server)")

    _h2(doc, "5.1 Clone and install")
    _code_block(
        doc,
        "git clone https://github.com/ronidas39/jira-mcp-server.git\n"
        "cd jira-mcp-server\n"
        "python -m venv .venv\n"
        "source .venv/bin/activate    # Windows: .venv\\Scripts\\activate\n"
        "pip install -e \".[dev]\"\n",
    )
    _para(
        doc,
        "The editable install (-e) means edits in src/ are picked up "
        "without reinstalling. The [dev] extra adds pytest, ruff, mypy, "
        "respx, and mongomock-motor.",
    )

    _h2(doc, "5.2 Environment configuration")
    _para(doc, "Copy the template and fill it in:")
    _code_block(doc, "cp .env.example .env\n")
    _para(
        doc,
        "Every setting in .env.example is documented inline. The minimum "
        "required for api_token mode:",
    )
    _kv_table(
        doc,
        ["Variable", "Required", "Example"],
        [
            ["JIRA_BASE_URL", "yes", "https://your-domain.atlassian.net"],
            ["JIRA_AUTH_MODE", "yes", "api_token"],
            ["JIRA_EMAIL", "yes (api_token)", "you@example.com"],
            ["JIRA_API_TOKEN", "yes (api_token)", "(opaque string)"],
            ["MONGO_URI", "yes", "mongodb://localhost:27017"],
            ["MONGO_DB", "no", "jira_mcp (default)"],
            ["MCP_TRANSPORT", "no", "stdio (default) or http"],
            ["MCP_HTTP_HOST", "no", "127.0.0.1 (default)"],
            ["MCP_HTTP_PORT", "no", "8765 (default)"],
            ["MCP_CORS_ORIGINS", "no", "http://localhost:3000 (default)"],
            ["LOG_LEVEL", "no", "INFO (default)"],
            ["ALLOW_DELETE_ISSUES", "no", "false (default; opt in for delete tool)"],
            ["JIRA_MAX_CONCURRENCY", "no", "8 (default httpx connection limit)"],
            ["JIRA_MAX_RETRIES", "no", "3 (default for 429 / 5xx backoff)"],
        ],
    )

    _h2(doc, "5.3 Jira credentials: API token mode (recommended for v1)")
    _para(
        doc,
        "API token auth is the fastest path. Each token is per-user and "
        "rotates independently of your Atlassian password.",
    )
    _bullet(doc, "Sign in to https://id.atlassian.com/manage-profile/security/api-tokens.")
    _bullet(doc, "Click Create API token; label it (for example jira-mcp-server).")
    _bullet(doc, "Copy the value; Atlassian only shows it once.")
    _bullet(
        doc,
        "Set JIRA_EMAIL to the Atlassian account that owns the token. The "
        "email must match exactly; a typo produces 401 at startup.",
    )
    _bullet(doc, "Set JIRA_API_TOKEN to the value you copied.")
    _para(doc, "Verify the credentials before going further:")
    _code_block(doc, "PYTHONPATH=src python scripts/smoke_jira.py\n")
    _para(
        doc,
        "The smoke script exercises every read-only path against your real "
        "Jira (myself, list_projects, list_users, resolve, list_custom_fields, "
        "search via /rest/api/3/search/jql, get_issue, list_boards, "
        "list_sprints, get_sprint). All steps must report ok.",
    )
    _para(doc, "To also run the write+cleanup loop (creates one ephemeral issue, comments, transitions, then deletes):")
    _code_block(doc, "ALLOW_DELETE_ISSUES=true PYTHONPATH=src python scripts/smoke_jira.py --include-writes\n")

    _h2(doc, "5.4 Jira credentials: OAuth 2.0 (3LO) mode")
    _para(
        doc,
        "OAuth is the right choice when you want per-user attribution in "
        "the audit log, when token rotation is policy-driven, or when you "
        "are deploying for multiple Atlassian accounts.",
    )
    _bullet(doc, "Create an OAuth 2.0 (3LO) app at https://developer.atlassian.com/console/myapps/.")
    _bullet(doc, "Set the callback URL on the Atlassian side to JIRA_OAUTH_REDIRECT_URI (default: http://localhost:9000/callback).")
    _bullet(doc, "Add the scopes: read:jira-work, write:jira-work, read:jira-user, offline_access.")
    _bullet(doc, "Copy the Client ID and Client Secret into JIRA_OAUTH_CLIENT_ID and JIRA_OAUTH_CLIENT_SECRET in .env.")
    _bullet(doc, "Set JIRA_AUTH_MODE=oauth.")
    _para(doc, "Run the one-shot login flow once:")
    _code_block(doc, "PYTHONPATH=src python scripts/oauth_login.py\n")
    _para(
        doc,
        "The script boots a local listener on the redirect-URI port, opens "
        "Atlassian's authorize page in your browser, captures the callback, "
        "exchanges the code, fetches the tenant cloud_id, and persists the "
        "token to MongoDB (collection oauth_tokens). After this completes "
        "the server picks the token up at startup and refreshes it eagerly "
        "when it is within sixty seconds of expiry.",
    )
    _para(
        doc,
        "If you see invalid_grant on a refresh, your refresh token is "
        "themselves invalid; rerun the login script.",
    )

    _h2(doc, "5.5 MongoDB setup")
    _para(doc, "Two paths.")
    _h3(doc, "Local Docker container")
    _code_block(doc, "docker compose up -d mongodb\n")
    _para(doc, "Then set MONGO_URI=mongodb://localhost:27017 in .env.")
    _h3(doc, "MongoDB Atlas (managed cloud)")
    _para(
        doc,
        "Use the SRV connection string Atlas provides. The connection "
        "helper passes certifi's CA bundle automatically so Atlas TLS "
        "works on macOS Python.org builds without any extra plumbing.",
    )
    _code_block(doc, "MONGO_URI=mongodb+srv://user:pass@cluster.mongodb.net/jira_mcp?retryWrites=true&w=majority\n")

    _h2(doc, "5.6 Initialize the database")
    _para(
        doc,
        "Idempotent. Creates the audit_log collection with three indexes "
        "(ts, tool+ts, actor+ts), the cache collection with a TTL index "
        "on expires_at, and the oauth_tokens collection (used only when "
        "OAuth mode is on).",
    )
    _code_block(doc, "PYTHONPATH=src python scripts/init_db.py\n")

    _h2(doc, "5.7 Run the server in stdio mode")
    _para(
        doc,
        "Stdio is the canonical transport for IDE assistants. The process "
        "reads JSON-RPC from stdin and writes to stdout; an MCP-capable "
        "client spawns the process and pipes the streams.",
    )
    _code_block(doc, "PYTHONPATH=src python -m jira_mcp\n")

    _h2(doc, "5.8 Run the server over streamable HTTP")
    _para(doc, "Required for the bundled web UI and for any remote MCP client.")
    _code_block(
        doc,
        "MCP_TRANSPORT=http MCP_HTTP_PORT=8765 PYTHONPATH=src python -m jira_mcp\n",
    )
    _para(
        doc,
        "The endpoint mounts at /mcp without a trailing slash. CORS is "
        "configured from MCP_CORS_ORIGINS (comma separated; the default "
        "permits http://localhost:3000 for the bundled UI). The server "
        "exposes Mcp-Session-Id as a response header so a browser client "
        "can reuse the same session across requests.",
    )

    _h2(doc, "5.9 Verify the HTTP server with curl")
    _para(doc, "Initialize a session and capture the Mcp-Session-Id:")
    _code_block(
        doc,
        "curl -i -X POST http://127.0.0.1:8765/mcp \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -H \"Accept: application/json, text/event-stream\" \\\n"
        "  -d '{\"jsonrpc\":\"2.0\",\"id\":1,\"method\":\"initialize\","
        "\"params\":{\"protocolVersion\":\"2025-03-26\",\"capabilities\":{},"
        "\"clientInfo\":{\"name\":\"curl\",\"version\":\"0\"}}}'\n",
    )
    _para(doc, "Then send the initialized notification (mandatory by spec) and list tools:")
    _code_block(
        doc,
        "SESSION=<value of Mcp-Session-Id from above>\n"
        "\n"
        "curl -X POST http://127.0.0.1:8765/mcp \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -H \"Accept: application/json, text/event-stream\" \\\n"
        "  -H \"Mcp-Session-Id: $SESSION\" \\\n"
        "  -d '{\"jsonrpc\":\"2.0\",\"method\":\"notifications/initialized\"}'\n"
        "\n"
        "curl -X POST http://127.0.0.1:8765/mcp \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -H \"Accept: application/json, text/event-stream\" \\\n"
        "  -H \"Mcp-Session-Id: $SESSION\" \\\n"
        "  -d '{\"jsonrpc\":\"2.0\",\"id\":2,\"method\":\"tools/list\"}'\n",
    )

    # 6. Frontend
    _h1(doc, "6. Frontend setup (Next.js dashboard)")

    _h2(doc, "6.1 Install dependencies")
    _code_block(doc, "cd web\nnpm install\n")
    _para(
        doc,
        "Stack: Next.js 15 with the App Router, TypeScript strict, Tailwind "
        "v4, shadcn/ui (latest), TanStack Query for client-side data, "
        "react-hook-form plus Zod for forms, the official "
        "@modelcontextprotocol/sdk for the API-route MCP client, and "
        "@anthropic-ai/sdk for the optional chat path.",
    )

    _h2(doc, "6.2 Environment")
    _code_block(doc, "cp .env.example .env.local\n")
    _kv_table(
        doc,
        ["Variable", "Purpose", "Default"],
        [
            ["MCP_SERVER_URL", "Streamable-HTTP endpoint of the Python MCP server", "http://localhost:8765/mcp"],
            ["ANTHROPIC_API_KEY", "Powers the /chat page; leave blank to disable", "(empty)"],
            ["JIRA_BROWSE_URL", "Used to build Open in Jira links", "https://ttmcp.atlassian.net"],
        ],
    )

    _h2(doc, "6.3 Run the dev server")
    _code_block(doc, "npm run dev\n")
    _para(doc, "The dashboard is at http://localhost:3000.")

    _h2(doc, "6.4 Production build")
    _code_block(doc, "npm run build\nnpm start\n")

    _h2(doc, "6.5 Quality gates")
    _code_block(doc, "npm run lint        # ESLint\nnpm run typecheck   # tsc --noEmit\nnpm run build       # Next.js production build\n")

    # 7. Docker
    _h1(doc, "7. Running with Docker Compose")
    _para(
        doc,
        "Two stack shapes are baked in. Both expect .env at the repo root "
        "to be populated.",
    )
    _h2(doc, "7.1 Headless: Mongo plus MCP server")
    _code_block(doc, "docker compose up\n")
    _para(
        doc,
        "Brings up MongoDB on 27017 and the MCP server on 8765 with the "
        "HTTP transport. Health checks gate the server on Mongo readiness.",
    )
    _h2(doc, "7.2 With UI: Mongo plus server plus dashboard")
    _code_block(doc, "docker compose --profile ui up\n")
    _para(
        doc,
        "Adds the web service on http://localhost:3000. The web container "
        "calls the server through the compose network at "
        "http://server:8765/mcp; the server's CORS allowlist is extended "
        "to permit that origin automatically.",
    )

    # 8. Using the frontend
    _h1(doc, "8. Using the dashboard")
    _para(
        doc,
        "Sidebar navigation on the left is permanent on desktop and "
        "collapses into a sheet on mobile. The dark mode toggle sits at "
        "the bottom of the sidebar. Every page is server-rendered against "
        "the live MCP server and degrades gracefully if a tool errors "
        "(an empty card with a hint instead of a thrown exception).",
    )

    _h2(doc, "8.1 Dashboard (/)")
    _para(
        doc,
        "Three at-a-glance cards built from list_projects, "
        "issues_by_status, workload_by_assignee, and a recent-activity "
        "JQL search (updated >= -2d ORDER BY updated DESC). Each card "
        "deep-links into the relevant detailed page.",
    )

    _h2(doc, "8.2 Issues (/issues)")
    _para(
        doc,
        "JQL input box (defaults to a project filter taken from the first "
        "project list_projects returns), a max-results selector, and a "
        "results table with key, summary, status, assignee, and last "
        "updated. Clicking a row navigates to the issue detail page. "
        "The Create button opens /issues/new.",
    )

    _h2(doc, "8.3 New issue (/issues/new)")
    _para(
        doc,
        "Form with project (select populated from list_projects), issue "
        "type (defaults to Task), summary, description (plain text; the "
        "server wraps it in ADF before sending to Jira), assignee account "
        "id, priority, and labels. The form is Zod-validated; a successful "
        "submit toasts the new key and redirects to /issues/{key}.",
    )

    _h2(doc, "8.4 Issue detail (/issues/{key})")
    _para(
        doc,
        "Three tabs: Overview shows the parsed Issue payload (status, "
        "assignee, reporter, labels, dates, project). Comments lists the "
        "expanded comment thread plus an Add comment form (calls "
        "add_comment). Transitions lists list_transitions output and "
        "renders one button per available transition; clicking a button "
        "calls transition_issue and re-fetches the issue.",
    )

    _h2(doc, "8.5 Projects (/projects, /projects/{key})")
    _para(
        doc,
        "List view shows every accessible project with key, name, and "
        "type. Project detail combines get_project (issue types, "
        "components, lead) with list_custom_fields filtered to the fields "
        "Jira reports as in-use for that project.",
    )

    _h2(doc, "8.6 Sprints (/sprints)")
    _para(
        doc,
        "Three sequential pickers: board (from list_boards), sprint state "
        "(active / closed / future), and finally a specific sprint. The "
        "Sprint Report card on the right shows committed, delivered, and "
        "at-risk counts from sprint_report. Caveats about how committed is "
        "approximated are documented inline because the agile API does "
        "not separate committed-at-start from currently-in-sprint without "
        "a full changelog scan.",
    )

    _h2(doc, "8.7 Analytics (/analytics)")
    _para(
        doc,
        "Project picker plus three cards: workload by assignee (a bar list "
        "with story-point totals when the field is configured), issues by "
        "status (a horizontally-stacked status histogram), and stale "
        "issues (a table with key, summary, days since update, and a link "
        "to /issues/{key}). The default stale window is thirty days and "
        "is adjustable in the UI.",
    )

    _h2(doc, "8.8 Chat (/chat)")
    _para(
        doc,
        "If ANTHROPIC_API_KEY is set in web/.env.local, the chat page "
        "becomes a streaming chat interface where Claude has the MCP "
        "server attached as tools and can call them on your behalf. The "
        "system prompt instructs the model to ask before any destructive "
        "write. If the key is unset, the page renders a placeholder with "
        "instructions to set it.",
    )

    _h2(doc, "8.9 Settings (/settings)")
    _para(
        doc,
        "Read-only inspection of MCP_SERVER_URL, JIRA_BROWSE_URL, and "
        "(redacted) ANTHROPIC_API_KEY status, plus a Test connection "
        "button that pings /api/mcp/tools to confirm the server is "
        "reachable and reports the count of tools it returned.",
    )

    # 9. Tools reference
    _h1(doc, "9. MCP tool reference")
    _para(
        doc,
        "The server registers twenty four tools. Tools that mutate Jira "
        "state additionally write one row to the audit_log collection.",
    )
    _kv_table(
        doc,
        ["Tool", "Purpose", "Audit"],
        [
            ["search_issues", "JQL search; returns a page of issue summaries.", "no"],
            ["get_issue", "Fetch one issue by key with optional comments and transitions expansion.", "no"],
            ["create_issue", "Create one issue; description is plain text wrapped in ADF.", "yes"],
            ["update_issue", "Partial update; only fields that are set are sent.", "yes"],
            ["transition_issue", "Move an issue through one workflow transition; optional comment.", "yes"],
            ["bulk_create_issues", "Up to fifty issues per call; per-item fallback on rejection.", "yes"],
            ["add_comment", "Append an ADF-wrapped comment to an issue.", "yes"],
            ["link_issues", "Create an issue link of a given type (Blocks, Relates, Duplicate).", "yes"],
            ["list_transitions", "List the workflow transitions available right now for an issue.", "no"],
            ["delete_issue", "Disabled by default; gated by ALLOW_DELETE_ISSUES.", "yes"],
            ["list_projects", "List every project the caller can read.", "no"],
            ["get_project", "Project metadata including issue types and components.", "no"],
            ["list_custom_fields", "Custom fields with their schema; filters to in-use ones.", "no"],
            ["list_users", "Search users; assignable variant when project_key is given.", "no"],
            ["resolve_user", "Resolve an email or display name to an accountId.", "no"],
            ["list_boards", "List agile boards; optional project filter.", "no"],
            ["list_sprints", "List sprints for a board with state filter.", "no"],
            ["get_sprint", "Single sprint details.", "no"],
            ["move_to_sprint", "Move issues into a sprint; batched at fifty.", "yes"],
            ["sprint_report", "Committed, delivered, and at-risk counts.", "no"],
            ["workload_by_assignee", "Per-assignee count plus story points (configurable field).", "no"],
            ["issues_by_status", "Status histogram for a project.", "no"],
            ["velocity", "Committed plus delivered points across the last N closed sprints.", "no"],
            ["stale_issues", "Issues not updated in N days; default thirty.", "no"],
        ],
    )

    # 10. Resources
    _h1(doc, "10. MCP resources reference")
    _kv_table(
        doc,
        ["URI", "Returns"],
        [
            ["jira://issue/{key}", "Full issue parsed from /rest/api/3/issue with renderedFields, transitions, changelog expansions."],
            ["jira://sprint/{id}", "Full sprint document from /rest/agile/1.0/sprint/{id}."],
            ["jira://project/{key}", "Project metadata with issueTypes."],
            ["jira://search?jql=...", "JQL result as raw search response (max fifty)."],
        ],
    )

    # 11. Prompts
    _h1(doc, "11. MCP prompts reference")
    _kv_table(
        doc,
        ["Slash command", "Required arguments", "Behaviour"],
        [
            ["/sprint-review", "board_id (sprint_id optional)", "Calls sprint_report and workload_by_assignee, formats a markdown retrospective."],
            ["/backlog-grooming", "project_key", "Runs stale_issues for thirty days, groups by status and priority, asks for confirmation before any write."],
            ["/triage-bugs", "project_key", "Searches issuetype=Bug AND status=Open, groups by priority, suggests priority changes."],
            ["/standup-summary", "project_key", "Searches updated>=-1d, groups by assignee, produces a markdown digest."],
        ],
    )

    # 12. Quality gates
    _h1(doc, "12. Quality gates and tests")
    _h2(doc, "12.1 Python")
    _code_block(
        doc,
        "ruff check src tests\n"
        "ruff format --check src tests\n"
        "mypy src                                    # strict\n"
        "PYTHONPATH=src python -m pytest -q --no-cov\n"
        "PYTHONPATH=src python scripts/smoke_jira.py # live read smoke\n",
    )
    _h2(doc, "12.2 Web")
    _code_block(doc, "cd web\nnpm run lint\nnpm run typecheck\nnpm run build\n")
    _h2(doc, "12.3 CI")
    _para(
        doc,
        "GitHub Actions runs every gate on push and pull request to main. "
        "See .github/workflows/ci.yml.",
    )

    # 13. Troubleshooting
    _h1(doc, "13. Troubleshooting")
    _kv_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["401 Unauthorized at startup", "Bad token or wrong email", "Regenerate the token; check the email matches the Atlassian account."],
            ["410 Gone on search", "Old endpoint", "Update the client; the server now uses /rest/api/3/search/jql."],
            ["403 Forbidden on writes", "Project permissions", "Make sure the user has the right project role."],
            ["Stream of 429 responses", "Rate-limited", "Lower JIRA_MAX_CONCURRENCY; the retry layer handles transient cases."],
            ["SSL: CERTIFICATE_VERIFY_FAILED", "Corp proxy with custom CA", "Set SSL_CERT_FILE to the CA bundle."],
            ["UI shows Connection failed", "MCP server not running", "Start it with MCP_TRANSPORT=http; check MCP_SERVER_URL."],
            ["CORS blocked in browser", "Origin missing from MCP_CORS_ORIGINS", "Add the UI origin (comma-separated) and restart the server."],
            ["Chat page shows placeholder", "ANTHROPIC_API_KEY missing", "Set it in web/.env.local and restart the dev server."],
            ["docker build fails on web/", "node_modules copied in", "Make sure web/.dockerignore is in place; run docker build with --no-cache."],
            ["mongo.connectivity.failed", "Mongo not reachable", "Start docker compose up -d mongodb; check MONGO_URI."],
            ["Empty audit_log on writes", "Tool ran but Mongo down", "Check .audit-buffer.jsonl; AuditRepository.flush_buffer replays it."],
        ],
    )

    # 14. Project structure
    _h1(doc, "14. Project structure")
    _code_block(
        doc,
        "jira-mcp-server/\n"
        "  src/jira_mcp/\n"
        "    __main__.py        process entry, lifecycle\n"
        "    auth/              ApiTokenAuth, OAuthProvider, AuthProvider Protocol\n"
        "    clients/           JiraClient (HTTP) plus IssueClient, ProjectClient,\n"
        "                       UserClient, SprintClient (domain helpers)\n"
        "    config/            pydantic-settings (HTTPS-only validator)\n"
        "    db/                MongoConnection, AuditRepository (append-only),\n"
        "                       CacheRepository (TTL), TokenRepository (OAuth)\n"
        "    models/            Pydantic entity models plus tool I/O schemas\n"
        "    prompts/           four slash-command templates\n"
        "    resources/         jira:// URI handlers\n"
        "    server/            create_app, lifespan, transport (stdio + http)\n"
        "    tools/             one module per group; register_all composes them\n"
        "    utils/             logging, errors, retry, JQL helpers, correlation\n"
        "  scripts/             init_db.py, smoke_jira.py, oauth_login.py,\n"
        "                       quality.sh, build_usage_docx.py\n"
        "  tests/               unit + fixtures + integration scaffolding\n"
        "  docs/                architecture, design, use cases, runbooks, this guide\n"
        "  web/                 Next.js plus TS plus Tailwind v4 plus shadcn UI\n"
        "  Dockerfile, docker-compose.yml, .github/workflows/ci.yml, pyproject.toml\n",
    )

    # 15. Pending / roadmap
    _h1(doc, "15. Pending and roadmap")
    _bullet(
        doc,
        "Live docker build verification: the Dockerfiles are written and "
        "ride the same wheel-build path that worked at project init, but I "
        "could not exercise docker build because the Docker daemon was not "
        "running on the development host. CI runs npm run build for the "
        "web side; add a docker build job to ci.yml when you want it.",
    )
    _bullet(
        doc,
        "Web unit tests: the dashboard has end-to-end paths only via the "
        "live MCP server. Adding Vitest plus Testing Library coverage on "
        "components/ would close the loop.",
    )
    _bullet(
        doc,
        "OAuth refresh against a real Atlassian app: the unit tests cover "
        "happy path and invalid_grant; running through the full 3LO flow "
        "with a developer.atlassian.com app is still recommended before "
        "production.",
    )
    _bullet(
        doc,
        "Demo recording: a short walkthrough video would be a useful "
        "addition for portfolio and classroom contexts.",
    )
    _bullet(
        doc,
        "Story-points override: analytics defaults to customfield_10016. "
        "Tenants that remapped the field need a settings entry; the hook "
        "(_resolve_story_points_field) is in place.",
    )
    _bullet(
        doc,
        "Webhook ingestion: out of scope for v1. A push-notification "
        "listener is the natural v2 follow on.",
    )

    # 16. Common operator workflows
    _h1(doc, "16. Common operator workflows")

    _h2(doc, "16.1 Triage a bug from a stack trace")
    _bullet(doc, "Open /issues/new.")
    _bullet(doc, "Pick the project and Bug as the issue type.")
    _bullet(doc, "Paste the trace into the description; the server wraps it in ADF.")
    _bullet(doc, "Submit; the toast links to the new key.")

    _h2(doc, "16.2 Bulk-create from meeting notes")
    _bullet(doc, "Open /chat (requires ANTHROPIC_API_KEY).")
    _bullet(doc, "Paste the notes and ask the model to bulk-create issues.")
    _bullet(doc, "The model calls bulk_create_issues; review the response and confirm before any further writes.")

    _h2(doc, "16.3 Sprint review")
    _bullet(doc, "Open /sprints, pick the board and the active sprint.")
    _bullet(doc, "The Sprint Report card shows committed, delivered, and at-risk counts.")
    _bullet(doc, "For a deeper write-up, open /chat and run /sprint-review with the same board id.")

    _h2(doc, "16.4 Stale-issue cleanup")
    _bullet(doc, "Open /analytics, pick the project.")
    _bullet(doc, "Stale issues card lists items not updated in thirty days.")
    _bullet(doc, "Either click into each one and update labels, or run /backlog-grooming in /chat for a guided pass.")

    _h2(doc, "16.5 Workload check before planning")
    _bullet(doc, "Open /analytics; the workload card lists per-assignee count plus points.")
    _bullet(doc, "Compare against the team's intended distribution before adding more work.")

    # 17. Screenshots gallery
    _h1(doc, "17. Screenshots gallery")
    _para(
        doc,
        "Full-page captures of every dashboard route live under "
        "docs/screenshots/. They sit alongside this guide so a reader can "
        "see what the UI is supposed to look like without running the "
        "stack first. Each PNG maps one-to-one with a route in section 8:",
    )
    _kv_table(
        doc,
        ["Route", "File"],
        [
            ["/", "docs/screenshots/dashboard.png"],
            ["/issues", "docs/screenshots/issues-list.png"],
            ["/issues/new", "docs/screenshots/issues-new.png"],
            ["/issues/{key}", "docs/screenshots/issue-detail.png"],
            ["/projects", "docs/screenshots/projects-list.png"],
            ["/projects/{key}", "docs/screenshots/project-detail.png"],
            ["/sprints", "docs/screenshots/sprints.png"],
            ["/analytics", "docs/screenshots/analytics.png"],
            ["/chat", "docs/screenshots/chat.png"],
            ["/settings", "docs/screenshots/settings.png"],
        ],
    )
    _para(
        doc,
        "Re-capture from the running stack with:",
    )
    _code_block(
        doc,
        "pip install playwright && playwright install chromium\n"
        "docker-compose --profile ui up -d   # or 'cd web && npm run dev'\n"
        "python scripts/capture_screenshots.py\n",
    )

    # 18. Re-running this guide
    _h1(doc, "18. Re-running this guide")
    _para(
        doc,
        "Re-generate this DOCX from source after substantive changes:",
    )
    _code_block(doc, "PYTHONPATH=src python scripts/build_usage_docx.py\n")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    build()
